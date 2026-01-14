import uuid
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.settings import settings
from app.models.transcriber import FastWhisperTranscriber


class TranscriberService:
    def __init__(self):
        self.transcriber: Optional[FastWhisperTranscriber] = None

    def init(self, **kwargs):
        self.transcriber = FastWhisperTranscriber(**kwargs)

    def is_ready(self) -> bool:
        return self.transcriber is not None

    def get(self) -> FastWhisperTranscriber:
        if not self.transcriber:
            raise RuntimeError("Transcriber is not initialized")
        return self.transcriber
    
    async def prepare_audio(
        self, 
        raw_bytes: bytes, 
        filename: str, 
        save_file: bool
    ):
        original = Path(filename)

        stem = original.stem[:20]
        suffix = original.suffix.lower()

        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        uid = uuid.uuid4().hex

        safe_name = f"{stem}_{timestamp}_{uid}{suffix}"

        if save_file:
            path = settings.AUDIO_DIR / safe_name
            path.write_bytes(raw_bytes)
            return path

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        tmp.write(raw_bytes)
        tmp.flush()
        return Path(tmp.name)
    
    async def export_result(
        self, 
        result: dict, 
        source_filename: str
    ) -> Path:
        stem = Path(source_filename).stem[:20]
        uid = uuid.uuid4().hex
        timestamp = datetime.now().strftime("%d_%m_%Y_%H%M%S")

        filename = f"{stem}_{timestamp}.docx"
        path = settings.TRANSCRIBE_RESULTS_DIR / filename

        document = Document()

        heading = document.add_heading("Результат транскрипции", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = document.add_paragraph(result.get("text", ""))

        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p_format = paragraph.paragraph_format
        p_format.line_spacing = 1.25
        p_format.space_after = Pt(6)
        p_format.space_before = Pt(6)

        document.save(path)
        return path


transcriber_service = TranscriberService()
