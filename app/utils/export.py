from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ExportFormat = Literal["docx", "txt", "md"]

def export_docx(result: dict, path: Path) -> Path:
    document = Document()

    heading = document.add_heading("Результат транскрипции", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.add_paragraph(result.get("text", ""))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = paragraph.paragraph_format
    p.line_spacing = 1.25
    p.space_after = Pt(6)
    p.space_before = Pt(6)

    document.save(path)
    return path


def export_txt(result: dict, path: Path) -> Path:
    path.write_text(result.get("text", ""), encoding="utf-8")
    return path


def export_markdown(result: dict, path: Path) -> Path:
    content = f"# Результат транскрипции\n\n{result.get('text', '')}"
    path.write_text(content, encoding="utf-8")
    return path


def export_result(result: dict, path: Path, format: ExportFormat) -> Path:
    match format:
        case "docx":
            return export_docx(result, path)
        case "txt":
            return export_txt(result, path)
        case "md":
            return export_markdown(result, path)
        case _:
            raise ValueError(f"Unsupported export format: {format}")
